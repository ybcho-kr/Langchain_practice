"""
문서 처리 프로세스
"""

import os
import re
import hashlib
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass, field
import markdown
from docx import Document
from unstructured.partition.auto import partition
from unstructured.chunking.title import chunk_by_title
from unstructured.staging.base import elements_to_json
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from transformers import AutoTokenizer

from src.utils.logger import get_logger, log_info, log_error
from src.utils.config import get_data_config, get_embedding_config
# semantic_chunker는 함수 내부에서 지연 import (순환 import 방지)


@dataclass
class DocumentMetadata:
    """문서 수준 메타데이터"""

    doc_id: str
    file_path: str
    version: str
    content_hash: str
    tags: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            'doc_id': self.doc_id,
            'file_path': self.file_path,
            'version': self.version,
            'content_hash': self.content_hash,
            'tags': self.tags,
        }


@dataclass
class DocumentChunk:
    """문서 청크"""
    content: str
    metadata: Dict[str, Any]
    chunk_id: str
    source_file: str
    chunk_index: int
    doc_id: str
    section_id: Optional[str] = None
    chunk_type: str = "text"
    heading_path: List[str] = field(default_factory=list)
    page_start: Optional[int] = None
    page_end: Optional[int] = None
    language: Optional[str] = None
    domain: Optional[str] = None
    embedding_version: Optional[str] = None
    doc_metadata: Optional[DocumentMetadata] = None


class DocumentProcessor:
    """문서 처리기"""
    
    def __init__(self):
        self.logger = get_logger()
        self.config = get_data_config()
        self.supported_extensions = {'.md', '.docx', '.txt', '.pdf'}
        self.default_language = getattr(self.config, 'language', None) or 'unknown'
        self.default_domain = getattr(self.config, 'domain', None) or 'general'
        embedding_config = get_embedding_config()
        self.embedding_version = getattr(embedding_config, 'name', None) or getattr(embedding_config, 'model_path', '')
        self.current_document_metadata: Optional[DocumentMetadata] = None

        # 파일 해시 캐시 (중복 처리 방지)
        self.file_hashes = {}

    def _build_document_metadata(self, file_path: Path) -> DocumentMetadata:
        """문서 수준 메타데이터 생성"""
        content_hash = self._calculate_file_hash(file_path)
        version = datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
        tags = [file_path.suffix.lstrip('.')] if file_path.suffix else []
        doc_id_source = f"{file_path.resolve()}::{content_hash}"
        doc_id = hashlib.md5(doc_id_source.encode('utf-8')).hexdigest()

        return DocumentMetadata(
            doc_id=doc_id,
            file_path=str(file_path),
            version=version,
            content_hash=content_hash,
            tags=tags,
        )
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """파일 해시 계산"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                return hashlib.md5(content).hexdigest()
        except Exception as e:
            self.logger.error(f"파일 해시 계산 실패: {file_path}, 오류: {str(e)}")
            return ""
    
    def _is_file_changed(self, file_path: Path) -> bool:
        """파일이 변경되었는지 확인"""
        current_hash = self._calculate_file_hash(file_path)
        file_key = str(file_path)
        
        if file_key not in self.file_hashes:
            self.file_hashes[file_key] = current_hash
            return True
        
        if self.file_hashes[file_key] != current_hash:
            self.file_hashes[file_key] = current_hash
            return True

        return False

    def _build_section_id(self, heading_path: List[str]) -> Optional[str]:
        """헤딩 경로를 기반으로 섹션 ID 생성"""
        if not heading_path:
            return None

        heading_str = '/'.join([h for h in heading_path if h])
        if not heading_str:
            return None

        if self.current_document_metadata:
            return f"{self.current_document_metadata.doc_id}:{heading_str}"

        return heading_str

    def _create_document_chunk(
        self,
        content: str,
        file_path: Path,
        chunk_index: int,
        metadata: Dict[str, Any],
        *,
        heading_path: Optional[List[str]] = None,
        section_id: Optional[str] = None,
        chunk_type: str = 'text',
        page_start: Optional[int] = None,
        page_end: Optional[int] = None,
        language: Optional[str] = None,
        domain: Optional[str] = None,
        chunk_id_override: Optional[str] = None,
    ) -> DocumentChunk:
        """DocumentChunk 생성 헬퍼"""
        doc_meta = self.current_document_metadata or self._build_document_metadata(file_path)
        heading_path = [h for h in (heading_path or []) if h]
        resolved_section_id = section_id or self._build_section_id(heading_path)

        merged_metadata = {
            'file_name': file_path.name,
            'file_path': str(file_path),
            'file_size': file_path.stat().st_size,
            'chunk_size': len(content),
            'chunk_index': chunk_index,
            'chunk_type': chunk_type,
            'heading_path': heading_path,
            'section_id': resolved_section_id,
            'page_start': page_start,
            'page_end': page_end,
            'language': language or self.default_language,
            'domain': domain or self.default_domain,
            'embedding_version': self.embedding_version,
            'doc_id': doc_meta.doc_id,
            'document': doc_meta.to_dict(),
            **metadata,
        }

        # None 값 제거하여 payload 간결화
        merged_metadata = {k: v for k, v in merged_metadata.items() if v is not None}

        chunk_id = chunk_id_override or f"{doc_meta.doc_id}_{chunk_index}"

        return DocumentChunk(
            content=content,
            metadata=merged_metadata,
            chunk_id=chunk_id,
            source_file=str(file_path),
            chunk_index=chunk_index,
            doc_id=doc_meta.doc_id,
            section_id=resolved_section_id,
            chunk_type=chunk_type,
            heading_path=heading_path,
            page_start=page_start,
            page_end=page_end,
            language=language or self.default_language,
            domain=domain or self.default_domain,
            embedding_version=self.embedding_version,
            doc_metadata=doc_meta,
        )
    
    def process_file(self, file_path: Union[str, Path], force_process: bool = False) -> List[DocumentChunk]:
        """단일 파일 처리"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_extensions:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {file_path.suffix}")

        # 파일 변경 확인 (중복 처리 방지)
        if not force_process and not self._is_file_changed(file_path):
            self.logger.info(f"파일이 변경되지 않음, 건너뛰기: {file_path}")
            return []

        self.logger.info(f"파일 처리 시작: {file_path}")
        self.current_document_metadata = self._build_document_metadata(file_path)

        try:
            # 파일 형식에 따른 처리
            if file_path.suffix.lower() == '.md':
                # md 파일은 헤더 계층 기반 청킹 처리
                chunks = self._process_markdown_to_chunks(file_path)
                self.logger.info(f"파일 처리 완료: {file_path}, 청크 수: {len(chunks)}")
                return chunks
            elif file_path.suffix.lower() == '.docx':
                # DOCX 파일은 의미 기반 청킹 사용 여부 확인
                use_semantic_chunking = getattr(self.config, 'use_semantic_chunking_for_docx', False)
                
                if use_semantic_chunking:
                    chunks = self._process_docx_with_semantic_chunking(file_path)
                else:
                    content = self._process_docx(file_path)
                    chunks = self._chunk_content(content, file_path)
            elif file_path.suffix.lower() == '.txt':
                content = self._process_text(file_path)
                chunks = self._chunk_content(content, file_path)
            elif file_path.suffix.lower() == '.pdf':
                content = self._process_pdf(file_path)
                chunks = self._chunk_content(content, file_path)
            else:
                raise ValueError(f"지원하지 않는 파일 형식: {file_path.suffix}")
            
            self.logger.info(f"파일 처리 완료: {file_path}, 청크 수: {len(chunks)}")
            return chunks
            
        except Exception as e:
            self.logger.error(f"파일 처리 실패: {file_path}, 오류: {str(e)}")
            raise
    
    def _process_markdown_to_chunks(self, file_path: Path) -> List[DocumentChunk]:
        """Markdown 파일을 헤더 계층 기반으로 청킹.
        - 헤더 계층(H1/H2/H3) 메타데이터 추출
        - 섹션 내부의 마크다운 표를 탐지하여 구조화 텍스트로 변환
        - 표 문서는 표 전용 분할(_split_table_content) 사용, 일반 문서는 토크나이저 분할 사용
        - 각 청크에 섹션 제목 프리픽스 추가
        """
        content = self._read_file_with_encoding(file_path)

        # 섹션 분리: 헤더 라인(#, ##, ###) 기준으로 split
        content_with_prepended_newline = '\n' + content
        sections = re.split(r'\n(?=#+)', content_with_prepended_newline)

        current_heading: Optional[str] = None
        current_sub_heading: Optional[str] = None
        docs_from_file: List[LCDocument] = []

        for section in sections:
            section = section.strip()
            if not section:
                continue

            lines = section.split('\n')
            header_line = lines[0]
            body = '\n'.join(lines[1:]).strip()

            level = 0
            if header_line.startswith('###'):
                level = 3
            elif header_line.startswith('##'):
                level = 2
            elif header_line.startswith('#'):
                level = 1

            title = re.sub(r'^#+\s*[\d\.]*\s*', '', header_line).strip()

            if level == 1:
                current_heading = title
                current_sub_heading = None
            elif level == 2:
                current_sub_heading = title

            metadata: Dict[str, Any] = {
                'source': str(file_path),
                'heading': current_heading,
                'sub-heading': current_sub_heading,
                'sub-sub-heading': title if level == 3 else None,
                'subject': title
            }

            if body:
                # 섹션 본문에서 표와 일반 텍스트를 분리
                text_parts, table_blocks = self._split_markdown_text_and_tables(body)

                for text_part in text_parts:
                    if text_part.strip():
                        docs_from_file.append(LCDocument(page_content=text_part.strip(), metadata=metadata))

                for table_block in table_blocks:
                    # 표 문서에는 is_table_data 플래그 추가
                    table_meta = dict(metadata)
                    table_meta['is_table_data'] = True
                    docs_from_file.append(LCDocument(page_content=table_block, metadata=table_meta))

        # 토크나이저 및 텍스트 분할
        # 임베딩 모델 경로를 설정 파일에서 가져옴
        embedding_config = get_embedding_config()
        tokenizer_model_path = embedding_config.model_path
        
        # 모델 경로가 없으면 기본값 사용 (호환성 유지)
        if not tokenizer_model_path:
            self.logger.warning("임베딩 모델 경로가 설정되지 않았습니다. 토크나이저 로드를 건너뜁니다.")
            tokenizer = None
        else:
            # 로컬 경로인 경우 Path 객체로 변환하여 절대 경로 사용
            tokenizer_path = Path(tokenizer_model_path)
            if tokenizer_path.exists() and tokenizer_path.is_dir():
                # 로컬 디렉토리인 경우 절대 경로로 변환 (Windows 경로 정규화)
                tokenizer_model_path = str(tokenizer_path.resolve()).replace('\\', '/')
            
            try:
                # 로컬 파일 시스템 경로임을 명시
                tokenizer = AutoTokenizer.from_pretrained(
                    tokenizer_model_path,
                    local_files_only=False,
                    trust_remote_code=False
                )
                self.logger.info(f"토크나이저 로드 성공: {tokenizer_model_path}")
            except Exception as e:
                tokenizer = None
                self.logger.warning(f"토크나이저 로드 실패 ({tokenizer_model_path}): {str(e)}. 표 문서는 전용 분할, 텍스트는 문자 분할로 대체합니다.")

        # 표 문서와 일반 문서를 분리
        table_docs: List[LCDocument] = [d for d in docs_from_file if d.metadata.get('is_table_data')]
        text_docs: List[LCDocument] = [d for d in docs_from_file if not d.metadata.get('is_table_data')]

        chunks: List[DocumentChunk] = []
        chunk_index = 0

        # 1) 표 문서: 표 전용 분할 사용
        for tdoc in table_docs:
            table_content = tdoc.page_content
            table_title = self._extract_table_title_from_content(table_content)
            table_chunks = self._split_table_content(table_content, table_title, self.config.chunk_size, self.config.chunk_overlap)

            for i, chunk_text in enumerate(table_chunks):
                # 메타데이터
                md = {
                    'is_table_data': True,
                    'content_type': 'table',
                    'source': tdoc.metadata.get('source'),
                    'heading': tdoc.metadata.get('heading'),
                    'sub-heading': tdoc.metadata.get('sub-heading'),
                    'sub-sub-heading': tdoc.metadata.get('sub-sub-heading'),
                    'subject': tdoc.metadata.get('subject'),
                    'table_title': table_title,
                }

                heading_path = [
                    md.get('heading'),
                    md.get('sub-heading'),
                    md.get('sub-sub-heading'),
                ]

                chunks.append(
                    self._create_document_chunk(
                        content=chunk_text,
                        file_path=file_path,
                        chunk_index=chunk_index,
                        metadata=md,
                        heading_path=heading_path,
                        chunk_type='table',
                    )
                )
                chunk_index += 1

        # 2) 일반 문서: 토크나이저 분할 사용 (실패 시 문자 분할)
        if text_docs:
            if tokenizer is not None:
                text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
                    tokenizer=tokenizer, chunk_size=self.config.chunk_size, chunk_overlap=self.config.chunk_overlap
                )
            else:
                from langchain_text_splitters import CharacterTextSplitter
                text_splitter = CharacterTextSplitter(chunk_size=self.config.chunk_size, chunk_overlap=self.config.chunk_overlap)

            split_docs = text_splitter.split_documents(text_docs)

            for doc in split_docs:
                subject = doc.metadata.get('subject')
                prefixed_content = f"### '{subject}'에 대한 설명입니다.\n\n" + re.sub(r'(?<!\.)\n', ' ', doc.page_content)

                md = {
                    'is_table_data': False,
                    'content_type': 'text',
                    'source': doc.metadata.get('source'),
                    'heading': doc.metadata.get('heading'),
                    'sub-heading': doc.metadata.get('sub-heading'),
                    'sub-sub-heading': doc.metadata.get('sub-sub-heading'),
                    'subject': subject,
                }

                heading_path = [
                    md.get('heading'),
                    md.get('sub-heading'),
                    md.get('sub-sub-heading'),
                ]

                chunks.append(
                    self._create_document_chunk(
                        content=prefixed_content,
                        file_path=file_path,
                        chunk_index=chunk_index,
                        metadata=md,
                        heading_path=heading_path,
                        chunk_type='text',
                    )
                )
                chunk_index += 1

        return chunks
    
    def _find_markdown_table_end(self, lines: List[str], start_line: int) -> int:
        """마크다운 표의 끝 줄 찾기"""
        current_line = start_line + 1
        
        # 구분선 건너뛰기
        if current_line < len(lines) and '---' in lines[current_line]:
            current_line += 1
        
        # 데이터 행들 처리
        while current_line < len(lines):
            line = lines[current_line].strip()
            
            # 빈 줄이면 표 종료
            if not line:
                break
            
            # 표 행이 아니면 종료
            if not self._is_markdown_table_row(line):
                break
            
            current_line += 1
        
        return current_line - 1
    
    def _get_table_overlap_text(self, text: str, overlap_size: int) -> str:
        """표 데이터에서 오버랩 텍스트 추출"""
        if len(text) <= overlap_size:
            return text
        
        # 마지막 부분에서 데이터 행 경계 찾기
        overlap_text = text[-overlap_size:]
        
        # 첫 번째 완전한 데이터 행부터 시작
        lines = overlap_text.split('\n')
        for i, line in enumerate(lines):
            if line.startswith('데이터 행'):
                return '\n'.join(lines[i:])
        
        return overlap_text
    
    
    def _extract_data_rows(self, table_content: str) -> List[str]:
        """표에서 데이터 행들 추출"""
        import re
        
        # 데이터 행 패턴
        data_pattern = r'데이터 행 \d+: ([^\n]+)'
        data_matches = re.findall(data_pattern, table_content)
        
        return data_matches
    def _split_table_content(self, table_content: str, table_title: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """표 내용을 청크 크기에 맞게 분할 (표 제목 포함)"""
        chunks = []
        
        # 표 헤더 정보 추출
        header_info = self._extract_table_header_info(table_content)
        
        # 데이터 행들 추출
        data_rows = self._extract_data_rows(table_content)
        
        if not data_rows:
            # 데이터가 없으면 전체를 하나의 청크로
            return [table_content]
        
        current_chunk = ""
        current_chunk_size = len(header_info)  # 헤더 정보 크기부터 시작
        
        for i, row in enumerate(data_rows):
            row_text = f"데이터 행 {i+1}: {row}\n"
            row_size = len(row_text)
            
            # 현재 청크에 추가했을 때 크기 확인
            if current_chunk_size + row_size > chunk_size and current_chunk:
                # 현재 청크 완성
                chunk_with_header = header_info + "\n" + current_chunk.strip()
                chunks.append(chunk_with_header)
                
                # 오버랩 처리
                overlap_text = self._get_table_overlap_text(current_chunk, chunk_overlap)
                current_chunk = overlap_text + row_text if overlap_text else row_text
                current_chunk_size = len(header_info) + len(current_chunk)
            else:
                # 현재 청크에 추가
                current_chunk += row_text
                current_chunk_size += row_size
        
        # 마지막 청크 추가
        if current_chunk:
            chunk_with_header = header_info + "\n" + current_chunk.strip()
            chunks.append(chunk_with_header)
        
        return chunks
    
    def _split_markdown_text_and_tables(self, content: str) -> tuple:
        """마크다운 본문에서 표와 일반 텍스트를 분리하여 반환
        반환: (text_parts: List[str], table_blocks: List[str])
        표는 _structure_markdown_table_with_context 포맷으로 구조화
        확장: 섹션 제목((a), (b) 등) 다음에 오는 표도 인식
        """
        lines = content.split('\n')
        text_accum: List[str] = []
        text_parts: List[str] = []
        table_blocks: List[str] = []

        i = 0
        while i < len(lines):
            # 섹션 제목 확인
            if self._is_section_title(lines[i]):
                # 섹션 제목을 텍스트에 포함
                text_accum.append(lines[i])
                i += 1
                # 다음 줄이 빈 줄이면 건너뛰기
                if i < len(lines) and not lines[i].strip():
                    i += 1
                continue
            
            # 표 시작 확인
            if self._is_markdown_table_start(lines, i):
                # 앞서 쌓인 텍스트를 하나의 블록으로 저장
                if text_accum:
                    text_parts.append('\n'.join(text_accum).strip())
                    text_accum = []

                table_data = self._extract_markdown_table(lines, i)
                if table_data:
                    # 섹션 제목이 바로 위에 있으면 컨텍스트에 포함
                    context = self._get_markdown_table_context(lines, i)
                    
                    # 바로 앞 줄에 섹션 제목이 있으면 추가
                    section_title = None
                    for j in range(max(0, i - 3), i):
                        if self._is_section_title(lines[j]):
                            section_title = lines[j].strip()
                            break
                    
                    if section_title:
                        context['section_title'] = section_title
                        # 표 제목 후보에도 추가
                        if 'title_candidates' not in context:
                            context['title_candidates'] = []
                        context['title_candidates'].append(section_title)
                    
                    structured_table = self._structure_markdown_table_with_context(table_data, context)
                    table_blocks.append(structured_table)
                    i = table_data['end_line'] + 1
                    continue
            # 표가 아니면 텍스트 누적
            text_accum.append(lines[i])
            i += 1

        if text_accum:
            text_parts.append('\n'.join(text_accum).strip())

        # 빈 블록 제거
        text_parts = [t for t in text_parts if t]
        table_blocks = [t for t in table_blocks if t]
        return text_parts, table_blocks

    
    def _structure_markdown_table_with_context(self, table_data: Dict, context: Dict[str, str]) -> str:
        """마크다운 표를 주변 텍스트 분석을 통해 구조화된 형태로 변환"""
        structured_lines = []
        
        # 표 제목 생성 (주변 텍스트 분석 포함)
        table_title = self._extract_markdown_table_title_with_context(table_data['headers'], context)
        
        # 표 정보 헤더
        structured_lines.append(f"표 제목: {table_title}")
        structured_lines.append(f"표 구조: {len(table_data['headers'])}개 열 ({' | '.join(table_data['headers'])})")
        structured_lines.append(f"총 행 수: {len(table_data['rows'])}행")
        structured_lines.append("")
        
        # 데이터 행 처리
        for row_idx, row_data in enumerate(table_data['rows'], 1):
            row_items = []
            for i, value in enumerate(row_data):
                if i < len(table_data['headers']):
                    row_items.append(f"{table_data['headers'][i]}: {value}")
            structured_lines.append(f"데이터 행 {row_idx}: {' | '.join(row_items)}")
        
        return '\n'.join(structured_lines)
    

    
    def _extract_markdown_table_title_with_context(self, headers: List[str], context: Dict[str, str]) -> str:
        """마크다운 표 제목을 주변 텍스트 분석을 통해 추출
        확장: 섹션 제목도 제목 후보로 고려
        """
        # 기본 제목
        default_title = "표 데이터"
        
        # 우선순위 1: 명시적 표 제목 (예: "표 5-18, ...")
        import re
        table_title_line = context.get('table_title_line')
        title_candidates = context.get('title_candidates', [])
        preferred_patterns = [
            r'^표\s*\d+[\-\.]?\d*[,\.:]?\s*(.+)$',  # 표 5-18, 제목 / 표 5.18: 제목 / 표 5 제목
            r'^Table\s*\d+[\-\.]?\d*[,\.:]?\s*(.+)$'
        ]
        
        def extract_main_title(line: str) -> Optional[str]:
            for pat in preferred_patterns:
                m = re.match(pat, line.strip(), flags=re.IGNORECASE)
                if m:
                    title_text = m.group(0).strip()
                    # 번호 제거한 순수 제목도 확보
                    pure = re.sub(r'^표\s*\d+[\-\.]?\d*[,\.:]?\s*', '', title_text, flags=re.IGNORECASE).strip()
                    return title_text if pure == '' else title_text
            return None
        
        # 1) 바로 위 줄이 명시적 표 제목이면 그것을 사용
        if table_title_line:
            main_title = extract_main_title(table_title_line)
            if main_title:
                return main_title
        
        # 2) 후보군에서 명시적 표 제목 우선 선택
        for cand in title_candidates:
            main_title = extract_main_title(cand)
            if main_title:
                return main_title
        
        # 우선순위 2: 섹션 제목(예: (a), (b))은 보조 정보이므로 메인 표 제목이 없을 때만 사용
        section_title = context.get('section_title')
        if section_title:
            if headers:
                return f"{section_title} {headers[0]}"
            return section_title
        
        # 우선순위 3: 기타 후보군 중 최적 후보 선택
        if title_candidates:
            best_title = self._select_best_title(title_candidates, headers)
            if best_title:
                return best_title
        
        return default_title
    




    def _get_markdown_table_context(self, lines: List[str], table_start_line: int) -> Dict[str, str]:
        """마크다운 표 주변 텍스트를 분석하여 제목 후보 추출
        확장: 섹션 제목((a), (b) 등) 및 "표 X-Y," 형식도 제목 후보에 포함
        """
        context = {
            'before': '',
            'after': '',
            'title_candidates': [],
            'section_title': None,
            'table_title_line': None  # 표 바로 위 줄
        }
        
        # 표 바로 위 줄 확인 (최우선)
        if table_start_line > 0:
            immediate_prev_line = lines[table_start_line - 1].strip()
            if immediate_prev_line:
                context['table_title_line'] = immediate_prev_line
        
        # 표 앞의 텍스트 분석 (최대 10개 줄로 확장하여 섹션 제목도 포함)
        before_texts = []
        section_title = None
        for i in range(max(0, table_start_line - 10), table_start_line):
            line = lines[i].strip()
            if not line:
                continue
            
            # 섹션 제목 확인
            if self._is_section_title(line):
                section_title = line
                before_texts.append(line)
                continue
            
            # 표 헤더나 구분선이 아니면 추가
            if not self._is_markdown_table_header(line) and '---' not in line:
                before_texts.append(line)
        
        context['before'] = '\n'.join(before_texts)  # 줄 구분 유지
        if section_title:
            context['section_title'] = section_title
        
        # 표 뒤의 텍스트 분석 (최대 3개 줄)
        table_end_line = self._find_markdown_table_end(lines, table_start_line)
        after_texts = []
        for i in range(table_end_line + 1, min(len(lines), table_end_line + 4)):
            line = lines[i].strip()
            if line and not self._is_markdown_table_header(line) and '---' not in line:
                after_texts.append(line)
        
        context['after'] = ' '.join(after_texts)
        
        # 제목 후보 추출 (before_text를 줄 단위로 전달)
        context['title_candidates'] = self._extract_title_candidates(context['before'], context['after'])
        
        # 표 바로 위 줄을 최우선 제목 후보로 추가
        if context['table_title_line']:
            table_title_line = context['table_title_line']
            # 표 번호 포함 형식이면 그대로, 아니면 표 번호 제거 후 추가
            import re
            if re.match(r'표\s*\d+', table_title_line, re.IGNORECASE):
                # 표 번호 포함 전체 라인 추가
                if table_title_line not in context['title_candidates']:
                    context['title_candidates'].insert(0, table_title_line)
                # 표 번호 제거한 버전도 추가
                title_without_number = re.sub(r'^표\s*\d+[-.]?\d*[,.]?\s*', '', table_title_line, flags=re.IGNORECASE).strip()
                if title_without_number and title_without_number not in context['title_candidates']:
                    context['title_candidates'].insert(1, title_without_number)
            else:
                # 표 번호가 없으면 그대로 최우선 추가
                if table_title_line not in context['title_candidates']:
                    context['title_candidates'].insert(0, table_title_line)
        
        # 섹션 제목을 제목 후보에 추가 (표 바로 위 줄 다음)
        if section_title:
            # 섹션 제목에서 괄호 제거 후 추가
            section_name = section_title.replace('(', '').replace(')', '').strip()
            if section_title not in context['title_candidates']:
                # 표 바로 위 줄이 있으면 그 다음에, 없으면 첫 번째에
                insert_pos = 1 if context['table_title_line'] else 0
                context['title_candidates'].insert(insert_pos, section_title)
        
        return context
 
    def _extract_markdown_table(self, lines: List[str], start_line: int) -> Dict:
        """마크다운 표 데이터 추출"""
        table_data = {
            'headers': [],
            'rows': [],
            'start_line': start_line,
            'end_line': start_line
        }
        
        # 헤더 행 처리
        header_line = lines[start_line]
        headers = [cell.strip() for cell in header_line.split('|')[1:-1]]  # 첫 번째와 마지막 빈 요소 제거
        table_data['headers'] = headers
        
        # 구분선 필수 반영: 다음 줄이 구분선인 경우만 표로 인정
        if not (start_line + 1 < len(lines) and '---' in lines[start_line + 1]):
            return None
        start_line += 1
        
        # 데이터 행 처리
        current_line = start_line + 1
        while current_line < len(lines):
            line = lines[current_line].strip()
            
            # 빈 줄이면 표 종료
            if not line:
                break
            
            # 표 행이 아니면 종료
            if not self._is_markdown_table_row(line):
                break
            
            # 데이터 행 추출
            row_data = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(row_data) == len(headers):
                table_data['rows'].append(row_data)
            
            current_line += 1
        
        table_data['end_line'] = current_line - 1
        return table_data
    def _is_markdown_table_row(self, line: str) -> bool:
        """마크다운 표 행인지 확인"""
        return '|' in line and line.count('|') >= 3          

    def _extract_table_header_info(self, table_content: str) -> str:
        """표 헤더 정보 추출 (표 제목, 구조, 행 수)"""
        import re
        
        # 표 헤더 정보 패턴
        header_pattern = r'(표 제목: [^\n]+\n표 구조: [^\n]+\n총 행 수: \d+행)'
        header_match = re.search(header_pattern, table_content)
        
        if header_match:
            return header_match.group(1)
        
        return "표 제목: 표 데이터\n표 구조: 알 수 없음\n총 행 수: 0행"
            

    def _is_markdown_table_start(self, lines: List[str], index: int) -> bool:
        """표 시작인지 엄격하게 판별
        - 현재 줄: 헤더 형태 (파이프 포함, 최소 3개 컬럼)
        - 다음 줄: --- 구분선이 컬럼 수에 맞춰 존재
        - 확장 패턴: (a), (b) 같은 섹션 제목 바로 다음에 오는 표도 인식
        """
        import re
        if index < 0 or index >= len(lines) - 1:
            return False
        
        # 현재 줄이 빈 줄이거나 섹션 제목((a), (b), (1), (2) 등)이면 건너뛰기
        current_line = lines[index].strip()
        if not current_line or self._is_section_title(current_line):
            # 섹션 제목 다음 줄도 확인
            if index + 2 < len(lines):
                return self._is_markdown_table_start(lines, index + 1)
            return False
        
        header = lines[index]
        separator = lines[index + 1]
        
        # 헤더는 파이프 기반 컬럼이 3개 이상
        if not self._is_markdown_table_header(header):
            return False
        
        # 구분선 정규식: | --- | :---: | ---: 등을 컬럼 수에 맞춰 포함
        # 파이프 기준으로 나누고 각 셀에 -가 3개 이상 존재하는지 확인
        header_cells = [c.strip() for c in header.split('|') if c.strip()]
        sep_cells = [c.strip() for c in separator.split('|') if c.strip()]
        
        if len(sep_cells) < max(1, len(header_cells)):
            return False
        
        def is_dash_cell(cell: str) -> bool:
            return re.fullmatch(r':?-{3,}:?', cell) is not None
        
        dash_like_count = sum(1 for cell in sep_cells if is_dash_cell(cell))
        
        # 구분선 셀의 대다수가 --- 형태여야 표로 인정
        return dash_like_count >= max(2, len(header_cells) - 1)
    
    def _is_section_title(self, line: str) -> bool:
        """섹션 제목인지 확인 (예: (a), (b), (1), (2), 등)"""
        import re
        # (a), (b), (1), (2), (가), (나) 등의 패턴
        section_patterns = [
            r'^\([a-zA-Z]\)',  # (a), (b), (A), (B)
            r'^\([0-9]+\)',    # (1), (2), (10)
            r'^\([가-힣]\)',   # (가), (나), (다)
            r'^\([가-힣]+\)',  # (가나), (다라)
        ]
        line_stripped = line.strip()
        for pattern in section_patterns:
            if re.match(pattern, line_stripped):
                return True
        return False
    

    def _extract_table_title_from_content(self, table_content: str) -> str:
        """표 내용에서 표 제목 추출"""
        import re
        
        # 표 제목 패턴 찾기
        title_match = re.search(r'표 제목: ([^\n]+)', table_content)
        if title_match:
            return title_match.group(1).strip()
        
        return "표 데이터"

    def _read_file_with_encoding(self, file_path: Path) -> str:
        """다양한 인코딩으로 파일 읽기"""
        encodings = ['utf-8', 'euc-kr', 'cp949', 'utf-8-sig', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                self.logger.info(f"파일 읽기 성공: {file_path}, 인코딩: {encoding}")
                return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                self.logger.error(f"파일 읽기 실패: {file_path}, 인코딩: {encoding}, 오류: {str(e)}")
                continue
        
        # 모든 인코딩 실패 시 에러 발생
        raise UnicodeDecodeError("파일을 읽을 수 없습니다", b"", 0, 1, "모든 인코딩 시도 실패")

    def _process_markdown(self, file_path: Path) -> str:
        """Markdown 파일 처리 (표 구조 보존)"""
        content = self._read_file_with_encoding(file_path)
        
        # Markdown 구조를 그대로 유지 (HTML 변환하지 않음)
        return content
    
    def _process_docx(self, file_path: Path) -> str:
        """DOCX 파일 처리 (기본: 문단과 표를 순서대로 읽기)"""
        doc = Document(file_path)
        content = []
        
        # 모든 요소를 순서대로 읽기 (문단과 표)
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 문단
                paragraph = None
                for p in doc.paragraphs:
                    if p._element == element:
                        paragraph = p
                        break
                if paragraph and paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            elif element.tag.endswith('tbl'):  # 표
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                if table:
                    # 간단한 표 처리
                    table_content = self._extract_table_content(table, file_path)
                    if table_content:
                        content.append(table_content)
        
        return '\n'.join(content)
        
    def _get_table_context(self, elements_with_context: List[Dict], table_index: int) -> Dict[str, str]:
        """표 주변 텍스트를 분석하여 제목 후보 추출"""
        context = {
            'before': '',
            'after': '',
            'title_candidates': []
        }
        
        # 표 앞의 텍스트 분석 (최대 3개 문단)
        before_texts = []
        for i in range(max(0, table_index - 3), table_index):
            if elements_with_context[i]['type'] == 'paragraph':
                before_texts.append(elements_with_context[i]['content'])
        
        context['before'] = ' '.join(before_texts)
        
        # 표 뒤의 텍스트 분석 (최대 2개 문단)
        after_texts = []
        for i in range(table_index + 1, min(len(elements_with_context), table_index + 3)):
            if elements_with_context[i]['type'] == 'paragraph':
                after_texts.append(elements_with_context[i]['content'])
        
        context['after'] = ' '.join(after_texts)
        
        # 제목 후보 추출
        context['title_candidates'] = self._extract_title_candidates(context['before'], context['after'])
        
        return context
    def _extract_title_candidates(self, before_text: str, after_text: str) -> List[str]:
        """표 주변 텍스트에서 제목 후보 추출
        확장: "표 X-Y," 또는 "표 X." 같은 패턴도 인식
        """
        candidates = []
        import re
        
        # 패턴 1: "표 X" 또는 "Table X" 패턴 (기존)
        table_patterns = [
            r'표\s*(\d+)[:：]\s*([^\n]+)',
            r'Table\s*(\d+)[:：]\s*([^\n]+)',
            r'표\s*(\d+)\s*([^\n]+)',
            r'Table\s*(\d+)\s*([^\n]+)'
        ]
        
        for pattern in table_patterns:
            matches = re.findall(pattern, before_text, re.IGNORECASE)
            for match in matches:
                if len(match) == 2:
                    candidates.append(match[1].strip())
        
        # 패턴 1-1: "표 X-Y," 또는 "표 X-Y." 형식 (쉼표 또는 마침표 포함)
        extended_table_patterns = [
            r'표\s*(\d+)[-.]\s*(\d+)[,.]\s*([^\n]+)',  # 표 5-21, 또는 표 5.22.
            r'표\s*(\d+)[,.]\s*([^\n]+)',  # 표 5-21, 또는 표 5.22.
            r'Table\s*(\d+)[-.]\s*(\d+)[,.]\s*([^\n]+)',  # Table 5-21, 또는 Table 5.22.
            r'Table\s*(\d+)[,.]\s*([^\n]+)',  # Table 5-21, 또는 Table 5.22.
        ]
        
        for pattern in extended_table_patterns:
            matches = re.findall(pattern, before_text, re.IGNORECASE)
            for match in matches:
                # match가 튜플이므로 길이 확인
                if len(match) >= 2:
                    # 마지막 요소가 제목 부분
                    title_part = match[-1].strip()
                    if title_part and len(title_part) > 3:  # 최소 길이 체크
                        candidates.append(title_part)
        
        # 패턴 1-2: "표 X-Y" (쉼표/마침표 없는 형식도 지원)
        table_with_dash_patterns = [
            r'표\s*(\d+)[-.]\s*(\d+)\s*([^\n]+)',  # 표 5-21 다음 줄에 제목
            r'Table\s*(\d+)[-.]\s*(\d+)\s*([^\n]+)',  # Table 5-21 다음 줄에 제목
        ]
        
        for pattern in table_with_dash_patterns:
            matches = re.findall(pattern, before_text, re.IGNORECASE)
            for match in matches:
                if len(match) >= 3:
                    title_part = match[-1].strip()
                    if title_part and len(title_part) > 3:
                        candidates.append(title_part)
        
        # 패턴 1-3: 표 바로 위 줄에 있는 전체 라인 (표 번호 포함 또는 제목만)
        # before_text를 줄 단위로 분석하여 표 바로 위 줄 확인
        before_lines = before_text.split('\n')
        if before_lines:
            # 마지막 줄 (표 바로 위 줄) 확인
            last_line = before_lines[-1].strip()
            
            # "표 X-Y," 또는 "표 X." 패턴 확인
            if re.match(r'표\s*\d+[-.]?\d*[,.]?\s*.+', last_line, re.IGNORECASE):
                # 전체 라인을 제목으로 사용 (표 번호 포함)
                candidates.append(last_line)
                # 표 번호 제거 후 제목만 추출
                title_without_number = re.sub(r'^표\s*\d+[-.]?\d*[,.]?\s*', '', last_line, flags=re.IGNORECASE).strip()
                if title_without_number:
                    candidates.append(title_without_number)
            
            # 표 번호 없이 제목만 있는 경우 (표 바로 위 줄)
            elif last_line and len(last_line) >= 5 and len(last_line) <= 100:
                # 특수 문자(■ 등)가 포함되어 있으면 제거
                cleaned_line = re.sub(r'^[■□▶▷●○◆◇★☆]+', '', last_line).strip()
                if cleaned_line:
                    candidates.append(cleaned_line)
        
        # 패턴 2: "다음은", "아래는", "위의" 등으로 시작하는 문장
        context_patterns = [
            r'다음은\s*([^.]*?)[의의]?\s*표',
            r'아래는\s*([^.]*?)[의의]?\s*표',
            r'위의\s*([^.]*?)[의의]?\s*표',
            r'다음\s*표는\s*([^.]*?)',
            r'아래\s*표는\s*([^.]*?)'
        ]
        
        for pattern in context_patterns:
            matches = re.findall(pattern, before_text, re.IGNORECASE)
            for match in matches:
                if match.strip():
                    candidates.append(match.strip())
        
        # 패턴 3: 짧은 문장 (20자 이하)이 표 바로 앞에 있는 경우
        sentences = re.split(r'[.!?。！？]', before_text)
        for sentence in sentences[-2:]:  # 마지막 2개 문장만 확인
            sentence = sentence.strip()
            if 5 <= len(sentence) <= 100 and not sentence.endswith('다') and not sentence.endswith('음'):
                candidates.append(sentence)
        
        # 패턴 4: 표 뒤의 텍스트에서 "이 표는", "위 표는" 등
        after_patterns = [
            r'이\s*표는\s*([^.]*?)',
            r'위\s*표는\s*([^.]*?)',
            r'다음\s*표는\s*([^.]*?)'
        ]
        
        for pattern in after_patterns:
            matches = re.findall(pattern, after_text, re.IGNORECASE)
            for match in matches:
                if match.strip():
                    candidates.append(match.strip())
        
        # 중복 제거 및 길이 필터링 (길이 제한 완화: 100자까지)
        unique_candidates = []
        seen = set()
        for candidate in candidates:
            candidate = candidate.strip()
            # 특수 문자만 있는 경우 제외
            if candidate and candidate not in seen and len(candidate) >= 3 and len(candidate) <= 100:
                # 특수 문자 패턴 체크 (■ 같은 것만 있는 경우 제외)
                if not re.match(r'^[■□▶▷●○◆◇★☆\s]+$', candidate):
                    unique_candidates.append(candidate)
                    seen.add(candidate)
        
        return unique_candidates
    
    def _extract_table_content_with_context(self, table, file_path: Path, context: Dict[str, str]) -> str:
        """표 내용 추출 (주변 텍스트 분석 포함)"""
        table_data = []
        
        # 표 헤더 추출 (첫 번째 행을 헤더로 가정)
        headers = []
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    headers.append(cell_text)
        
        # 표 정보 헤더
        if headers:
            # 표 제목을 주변 텍스트에서 추출
            table_title = self._extract_table_title_with_context(table, headers, context, str(file_path))
            
            table_data.append(f"표 제목: {table_title}")
            table_data.append(f"표 구조: {len(headers)}개 열 ({' | '.join(headers)})")
            table_data.append(f"총 행 수: {len(table.rows)}행")
            table_data.append("")
        
        # 데이터 행 처리 (첫 번째 행은 헤더이므로 제외)
        for row_idx, row in enumerate(table.rows[1:], 1):  # 헤더 제외
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_data.append(cell_text)
            
            if row_data:
                # 헤더가 있으면 키-값 쌍으로 변환
                if headers and len(headers) == len(row_data):
                    row_items = []
                    for i, value in enumerate(row_data):
                        if i < len(headers):
                            row_items.append(f"{headers[i]}: {value}")
                    table_data.append(f"데이터 행 {row_idx}: {' | '.join(row_items)}")
                else:
                    # 헤더가 없거나 개수가 맞지 않으면 단순 나열
                    table_data.append(f"데이터 행 {row_idx}: {' | '.join(row_data)}")
        
        if table_data:
            # 실제 표 데이터만 포함 (메타데이터 제거)
            table_content = '\n'.join(table_data)
            return table_content
        
        return ""


    def _extract_table_content(self, table, file_path: Path) -> str:
        """표 내용 추출 (하나의 청크로 유지)"""
        table_data = []
        
        # 표 헤더 추출 (첫 번째 행을 헤더로 가정)
        headers = []
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    headers.append(cell_text)
        
        # 표 정보 헤더
        if headers:
            # 표 제목을 동적으로 추출
            table_title = self._extract_table_title_with_context(table, headers, str(file_path))
            
            table_data.append(f"표 제목: {table_title}")
            table_data.append(f"표 구조: {len(headers)}개 열 ({' | '.join(headers)})")
            table_data.append(f"총 행 수: {len(table.rows)}행")
            table_data.append("")
        
        # 데이터 행 처리 (첫 번째 행은 헤더이므로 제외)
        for row_idx, row in enumerate(table.rows[1:], 1):  # 헤더 제외
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_data.append(cell_text)
            
            if row_data:
                # 헤더가 있으면 키-값 쌍으로 변환
                if headers and len(headers) == len(row_data):
                    row_items = []
                    for i, value in enumerate(row_data):
                        if i < len(headers):
                            row_items.append(f"{headers[i]}: {value}")
                    table_data.append(f"데이터 행 {row_idx}: {' | '.join(row_items)}")
                else:
                    # 헤더가 없거나 개수가 맞지 않으면 단순 나열
                    table_data.append(f"데이터 행 {row_idx}: {' | '.join(row_data)}")
        
        if table_data:
            # 실제 표 데이터만 포함 (메타데이터 제거)
            table_content = '\n'.join(table_data)
            return table_content
        
        return ""
    def _extract_table_title_with_context(self, table, headers, context: Dict[str, str], file_path: str = None) -> str:
        """표 제목을 주변 텍스트 분석을 통해 추출"""
        # 기본 제목
        default_title = "표 데이터"
        
        # 방법 1: 주변 텍스트에서 제목 후보 추출
        title_candidates = context.get('title_candidates', [])
        
        if title_candidates:
            # 가장 적절한 제목 선택 (길이와 내용 기반)
            best_title = self._select_best_title(title_candidates, headers)
            if best_title:
                return best_title
        
        # 방법 2: 헤더 내용을 기반으로 제목 생성
        if headers:
            if len(headers) >= 2:
                return f"{headers[0]} 및 {headers[-1]}"
            else:
                return headers[0] if headers else default_title
        
        return default_title
        
    def _extract_table_title(self, table, headers, file_path: str = None) -> str:
        """표 제목을 동적으로 추출"""
        # 기본 제목
        default_title = "표 데이터"
        
        # 방법 0: 파일명 기반 표 제목 추출 (최우선)
        if file_path:
            file_name = Path(file_path).stem.lower()
            if '유중가스' in file_name and '결함' in file_name:
                return "유중가스 결함 유형 및 조치"
            elif 'dga' in file_name and 'fault' in file_name:
                return "유중가스 결함 유형 및 조치"
        
        # 방법 1: 첫 번째 셀이 의미있는 제목인지 확인
        if table.rows and len(table.rows) > 0:
            first_cell = table.rows[0].cells[0].text.strip() if table.rows[0].cells else ""
            
            # 첫 번째 셀이 헤더와 다른지 확인 (제목일 가능성)
            if first_cell and first_cell not in headers and len(first_cell) > 3:
                return first_cell
        
        # 방법 2: 헤더 내용을 기반으로 제목 생성
        if headers:
            # 일반적인 표 제목 패턴 매칭
            header_text = ' '.join(headers).lower()
            
            # 한국어 패턴 매칭 추가 (우선순위 순서)
            if '유중가스' in header_text:
                return "유중가스 결함 유형 및 조치"
            elif '결함' in header_text and ('유형' in header_text or 'type' in header_text):
                return "결함 유형 및 조치"
            elif 'fault' in header_text and 'type' in header_text:
                return "결함 유형 및 조치"
            elif '가스' in header_text and ('유형' in header_text or 'type' in header_text):
                return "가스 유형 분류"
            elif 'gas' in header_text and 'type' in header_text:
                return "가스 유형 분류"
            elif 'diagnosis' in header_text or '진단' in header_text:
                return "진단 기준"
            elif 'test' in header_text or '시험' in header_text:
                return "시험 항목"
            else:
                # 헤더의 첫 번째와 마지막을 조합
                if len(headers) >= 2:
                    return f"{headers[0]} 및 {headers[-1]}"
                else:
                    return headers[0] if headers else default_title
        
        return default_title


    def _is_markdown_table_header(self, line: str) -> bool:
        """마크다운 표 헤더인지 확인 (단일 행 기준 약식 검사)"""
        return '|' in line and line.count('|') >= 3

    def _select_best_title(self, candidates: List[str], headers: List[str]) -> str:
        """제목 후보 중 가장 적절한 제목 선택"""
        if not candidates:
            return None
        
        # 점수 기반 선택
        scored_candidates = []
        
        for candidate in candidates:
            score = 0
            
            # 길이 점수 (10-30자가 적절)
            if 10 <= len(candidate) <= 30:
                score += 3
            elif 5 <= len(candidate) <= 50:
                score += 1
            
            # 헤더와의 연관성 점수
            if headers:
                for header in headers:
                    if any(word in candidate for word in header.split()):
                        score += 2
            
            # 키워드 점수 (표, 데이터, 정보 등)
            keywords = ['표', '데이터', '정보', '결과', '분석', '진단', '기준', '조치']
            for keyword in keywords:
                if keyword in candidate:
                    score += 1
            
            # 문장 완성도 점수 (마침표로 끝나지 않는 것)
            if not candidate.endswith('.') and not candidate.endswith('다'):
                score += 1
            
            scored_candidates.append((candidate, score))
        
        # 점수가 가장 높은 후보 선택
        if scored_candidates:
            best_candidate = max(scored_candidates, key=lambda x: x[1])
            if best_candidate[1] > 0:  # 최소 점수 이상
                return best_candidate[0]
        
        return None
       
    def _process_text(self, file_path: Path) -> str:
        """텍스트 파일 처리"""
        return self._read_file_with_encoding(file_path)
    
    def _process_pdf(self, file_path: Path) -> str:
        """PDF 파일 처리 (unstructured 사용)"""
        elements = partition(str(file_path))
        content = []
        
        for element in elements:
            if hasattr(element, 'text') and element.text.strip():
                content.append(element.text.strip())
        
        return '\n'.join(content)
    
    def _chunk_content(self, content: str, file_path: Path) -> List[DocumentChunk]:
        """개선된 한국어 청킹 로직 (표 데이터 우선 처리)"""
        chunks = []
        
        # 1단계: 표 데이터와 일반 텍스트 분리
        table_sections, text_sections = self._separate_table_and_text(content)
        
        chunk_index = 0
        
        # 2단계: 표 데이터 처리 (우선순위)
        for table_section in table_sections:
            # 연속된 표 병합 확인
            merged_table = self._merge_consecutive_tables(table_section)
            
            # 표 크기 확인 및 분할 처리
            if len(merged_table) > self.config.chunk_size:
                # 표가 너무 크면 행 단위로 분할
                table_chunks = self._split_large_table(merged_table, file_path, chunk_index)
                chunks.extend(table_chunks)
                chunk_index += len(table_chunks)
            else:
                # 표가 작으면 하나의 청크로 처리
                chunks.append(self._create_chunk(
                    merged_table,
                    file_path,
                    chunk_index
                ))
                chunk_index += 1
        
        # 3단계: 일반 텍스트 처리 (한국어 문장 기반 청킹)
        for text_section in text_sections:
            text_chunks = self._chunk_text_section(text_section, file_path, chunk_index)
            chunks.extend(text_chunks)
            chunk_index += len(text_chunks)
        
        return chunks
    
    def _separate_table_and_text(self, content: str) -> tuple:
        """표 데이터와 일반 텍스트를 분리"""
        import re
        
        # 표 데이터 패턴 찾기 (표 제목부터 다음 표 제목 또는 문서 끝까지)
        table_pattern = r'(표 제목: [^\n]+\n표 구조: [^\n]+\n총 행 수: \d+행\n\n.*?)(?=표 제목:|$)'
        table_matches = re.findall(table_pattern, content, re.DOTALL)
        
        # 표 데이터 추출
        table_sections = []
        remaining_content = content
        
        for table_match in table_matches:
            table_sections.append(table_match.strip())
            remaining_content = remaining_content.replace(table_match, "")
        
        # 일반 텍스트 추출 (더 큰 단위로 합치기)
        text_sections = self._merge_text_sections(remaining_content)
        
        return table_sections, text_sections
    
    def _merge_text_sections(self, content: str) -> List[str]:
        """텍스트 섹션을 효율적으로 합치기 (청크 수 최소화)"""
        # 먼저 문단 단위로 분리
        paragraphs = [section.strip() for section in content.split('\n\n') if section.strip()]
        
        merged_sections = []
        current_section = ""
        
        for paragraph in paragraphs:
            # 다음 섹션 크기 예상
            next_section_size = len(current_section) + len(paragraph) + (2 if current_section else 0)  # \n\n 포함
            
            # 청크 크기의 90%까지 허용하여 효율성 증대
            max_section_size = int(self.config.chunk_size * 0.9)
            
            if next_section_size > max_section_size and current_section:
                # 현재 섹션 완성
                merged_sections.append(current_section.strip())
                current_section = paragraph
            else:
                # 문단을 현재 섹션에 추가
                if current_section:
                    current_section += '\n\n' + paragraph
                else:
                    current_section = paragraph
        
        # 마지막 섹션 처리
        if current_section:
            merged_sections.append(current_section.strip())
        
        return merged_sections
    
    def _chunk_text_section(self, text: str, file_path: Path, start_index: int) -> List[DocumentChunk]:
        """일반 텍스트를 효율적으로 청킹 (문장 합치기 최적화)"""
        chunks = []
        
        sentences = self._split_korean_sentences(text)
        current_chunk = ""
        chunk_index = start_index
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # 다음 청크 크기 예상 (현재 청크 + 문장 + 공백)
            next_chunk_size = len(current_chunk) + len(sentence) + (1 if current_chunk else 0)
            
            # 청크 크기 확인 (최대 사이즈의 90%까지 허용하여 효율성 증대)
            max_chunk_size = int(self.config.chunk_size * 0.9)
            
            if next_chunk_size > max_chunk_size and current_chunk:
                # 현재 청크 완성
                chunks.append(self._create_chunk(
                    current_chunk.strip(),
                    file_path,
                    chunk_index
                ))
                chunk_index += 1
                
                # 오버랩 텍스트 추출 및 다음 청크 시작
                overlap_text = self._get_overlap_text(current_chunk)
                if overlap_text:
                    current_chunk = overlap_text + ' ' + sentence
                else:
                    current_chunk = sentence
            elif next_chunk_size > self.config.chunk_size and not current_chunk:
                # 단일 문장이 청크 크기를 초과하는 경우 강제 분할
                sub_chunks = self._split_large_text_with_overlap(sentence, file_path, chunk_index)
                chunks.extend(sub_chunks)
                chunk_index += len(sub_chunks)
            else:
                # 문장을 현재 청크에 추가
                if current_chunk:
                    current_chunk += ' ' + sentence
                else:
                    current_chunk = sentence
        
        # 마지막 청크 처리
        if current_chunk:
            chunks.append(self._create_chunk(
                current_chunk.strip(),
                file_path,
                chunk_index
            ))
        
        return chunks
    
    
    
    
    def _merge_consecutive_tables(self, content: str) -> str:
        """연속된 표 데이터를 병합"""
        import re
        
        # 표 데이터 패턴 찾기
        table_pattern = r'(표 제목: [^\n]+\n표 구조: [^\n]+\n총 행 수: \d+행[\s\S]*?)(?=\n\n|$)'
        table_matches = re.findall(table_pattern, content)
        
        if len(table_matches) < 2:
            return content
        
        # 연속된 표인지 확인 (헤더가 같은지)
        merged_content = content
        i = 0
        while i < len(table_matches) - 1:
            current_table = table_matches[i]
            next_table = table_matches[i + 1]
            
            # 헤더가 같은지 확인
            if self._has_same_headers(current_table, next_table):
                # 연속된 표로 인식하여 병합
                merged_table = self._merge_table_data(current_table, next_table)
                merged_content = merged_content.replace(current_table, merged_table)
                merged_content = merged_content.replace(next_table, "")
                i += 2  # 두 표를 건너뛰기
            else:
                i += 1
        
        return merged_content
    
    def _has_same_headers(self, table1: str, table2: str) -> bool:
        """두 표의 헤더가 같은지 확인"""
        import re
        
        # 표 구조에서 헤더 추출
        pattern1 = re.search(r'표 구조: \d+개 열 \(([^)]+)\)', table1)
        pattern2 = re.search(r'표 구조: \d+개 열 \(([^)]+)\)', table2)
        
        if pattern1 and pattern2:
            headers1 = pattern1.group(1).split(' | ')
            headers2 = pattern2.group(1).split(' | ')
            return headers1 == headers2
        
        return False
    
    def _merge_table_data(self, table1: str, table2: str) -> str:
        """두 표 데이터를 병합"""
        import re
        
        # 첫 번째 표의 행 수 추출
        rows_match1 = re.search(r'총 행 수: (\d+)행', table1)
        rows_match2 = re.search(r'총 행 수: (\d+)행', table2)
        
        if rows_match1 and rows_match2:
            rows1 = int(rows_match1.group(1))
            rows2 = int(rows_match2.group(1))
            total_rows = rows1 + rows2
            
            # 두 번째 표의 데이터 행들 추출
            data_rows2 = []
            lines2 = table2.split('\n')
            for line in lines2:
                if line.startswith('데이터 행'):
                    data_rows2.append(line)
            
            # 첫 번째 표에 두 번째 표의 데이터 행들 추가 (행 번호 재조정)
            merged_table = table1
            for i, data_row in enumerate(data_rows2):
                new_row_num = rows1 + i + 1
                new_data_row = re.sub(r'데이터 행 \d+:', f'데이터 행 {new_row_num}:', data_row)
                merged_table += '\n' + new_data_row
            
            # 총 행 수 업데이트
            merged_table = re.sub(r'총 행 수: \d+행', f'총 행 수: {total_rows}행', merged_table)
            
            return merged_table
        
        return table1
    
    def _split_large_table(self, table_content: str, file_path: Path, start_index: int) -> List[DocumentChunk]:
        """큰 표를 행 단위로 분할하여 연속된 청크 생성"""
        chunks = []
        lines = table_content.split('\n')
        
        # 표 헤더 정보 추출
        table_title = ""
        table_structure = ""
        total_rows = 0
        
        for line in lines:
            if line.startswith("표 제목:"):
                table_title = line
            elif line.startswith("표 구조:"):
                table_structure = line
            elif line.startswith("총 행 수:"):
                import re
                match = re.search(r'총 행 수: (\d+)행', line)
                if match:
                    total_rows = int(match.group(1))
                break
        
        # 데이터 행들 추출
        data_rows = []
        for line in lines:
            if line.startswith("데이터 행"):
                data_rows.append(line)
        
        # 행 단위로 청크 생성
        current_chunk_lines = [table_title, table_structure]
        current_rows = []
        chunk_index = start_index
        
        for i, data_row in enumerate(data_rows):
            # 현재 청크에 행 추가 시 크기 확인
            test_chunk = '\n'.join(current_chunk_lines + current_rows + [data_row])
            
            if len(test_chunk) > self.config.chunk_size and len(current_rows) > 0:
                # 현재 청크 완성
                current_rows_text = '\n'.join(current_rows)
                chunk_content = '\n'.join(current_chunk_lines) + f'\n총 행 수: {len(current_rows)}행\n\n{current_rows_text}'
                
                chunks.append(self._create_chunk(
                    chunk_content,
                    file_path,
                    chunk_index
                ))
                chunk_index += 1
                
                # 새 청크 시작 (헤더 포함)
                current_chunk_lines = [table_title, table_structure]
                current_rows = [data_row]
            else:
                current_rows.append(data_row)
        
        # 마지막 청크 처리
        if current_rows:
            current_rows_text = '\n'.join(current_rows)
            chunk_content = '\n'.join(current_chunk_lines) + f'\n총 행 수: {len(current_rows)}행\n\n{current_rows_text}'
            
            chunks.append(self._create_chunk(
                chunk_content,
                file_path,
                chunk_index
            ))
        
        return chunks
    
    def _get_overlap_text(self, text: str) -> str:
        if len(text) <= self.config.chunk_overlap:
            return text
        
        # 마지막 부분에서 문장 경계 찾기
        overlap_text = text[-self.config.chunk_overlap:]
        
        # 문장 경계에서 자르기 (첫 번째 문장부호 이후)
        sentences = self._split_korean_sentences(overlap_text)
        if len(sentences) > 1:
            # 첫 번째 문장을 제외하고 나머지 반환
            return ' '.join(sentences[1:])
        else:
            # 문장 경계가 없으면 그대로 반환
            return overlap_text
    
    def _split_korean_sentences(self, text: str) -> List[str]:
        """한국어 문장부호를 고려한 문장 분할"""
        import re
        
        # 한국어 문장부호 패턴 (마침표, 물음표, 느낌표, 줄임표 등)
        sentence_endings = r'[.!?。！？]|\.{2,}|…'
        
        # 문장부호로 분할하되, 문장부호는 유지
        sentences = re.split(f'({sentence_endings})', text)
        
        # 문장부호와 문장을 다시 결합
        result = []
        for i in range(0, len(sentences), 2):
            if i + 1 < len(sentences):
                sentence = sentences[i].strip() + sentences[i + 1]
                if sentence.strip():
                    result.append(sentence.strip())
            elif sentences[i].strip():
                result.append(sentences[i].strip())
        
        return result
    
    def _split_large_text_with_overlap(self, text: str, file_path: Path, start_index: int) -> List[DocumentChunk]:
        """큰 텍스트를 여러 청크로 분할 (한국어 어절 고려, 오버랩 포함)"""
        chunks = []
        words = self._split_korean_words(text)
        current_chunk = ""
        chunk_index = start_index
        overlap_buffer = ""
        
        for word in words:
            if len(current_chunk) + len(word) + 1 > self.config.chunk_size:
                if current_chunk:
                    chunks.append(self._create_chunk(
                        current_chunk.strip(),
                        file_path,
                        chunk_index
                    ))
                    chunk_index += 1
                    
                    # 오버랩 버퍼 설정
                    overlap_buffer = self._get_overlap_text(current_chunk)
                    current_chunk = overlap_buffer + ' ' + word if overlap_buffer else word
                else:
                    # 단일 단어가 청크 크기를 초과하는 경우 (매우 드문 경우)
                    chunks.append(self._create_chunk(
                        word,
                        file_path,
                        chunk_index
                    ))
                    chunk_index += 1
            else:
                if current_chunk:
                    current_chunk += ' ' + word
                else:
                    current_chunk = word
        
        # 마지막 청크 추가
        if current_chunk:
            chunks.append(self._create_chunk(
                current_chunk.strip(),
                file_path,
                chunk_index
            ))
        
        return chunks
    
    def _split_korean_words(self, text: str) -> List[str]:
        """한국어 어절을 고려한 단어 분할"""
        import re
        
        # 한국어 어절 패턴 (한글, 숫자, 영문, 특수문자 조합)
        # 한글 어절: 한글+숫자+영문+특수문자 조합
        korean_pattern = r'[가-힣]+[0-9a-zA-Z가-힣\s\-_\.]*[가-힣0-9a-zA-Z]'
        
        # 영문/숫자 패턴
        english_pattern = r'[a-zA-Z]+[0-9a-zA-Z\-_\.]*'
        number_pattern = r'[0-9]+[0-9\.\-_]*'
        
        # 특수문자 패턴
        special_pattern = r'[^\s가-힣a-zA-Z0-9]+'
        
        # 모든 패턴을 결합
        all_patterns = f'({korean_pattern}|{english_pattern}|{number_pattern}|{special_pattern})'
        
        # 패턴 매칭으로 단어 추출
        words = re.findall(all_patterns, text)
        
        # 공백으로도 분할하여 누락된 단어들 추가
        space_words = text.split()
        
        # 두 결과를 합치고 중복 제거 (순서 유지)
        result = []
        seen = set()
        for word in words + space_words:
            word = word.strip()
            if word and word not in seen:
                result.append(word)
                seen.add(word)
        
        return result
    
    def _split_large_text(self, text: str, file_path: Path, start_index: int) -> List[DocumentChunk]:
        """큰 텍스트를 여러 청크로 분할"""
        chunks = []
        words = text.split()
        current_chunk = ""
        chunk_index = start_index
        
        for word in words:
            if len(current_chunk) + len(word) + 1 > self.config.chunk_size:
                if current_chunk:
                    chunks.append(self._create_chunk(
                        current_chunk.strip(),
                        file_path,
                        chunk_index
                    ))
                    chunk_index += 1
                    current_chunk = word
                else:
                    # 단일 단어가 청크 크기를 초과하는 경우
                    chunks.append(self._create_chunk(
                        word,
                        file_path,
                        chunk_index
                    ))
                    chunk_index += 1
            else:
                if current_chunk:
                    current_chunk += ' ' + word
                else:
                    current_chunk = word
        
        # 마지막 청크 추가
        if current_chunk:
            chunks.append(self._create_chunk(
                current_chunk.strip(),
                file_path,
                chunk_index
            ))
        
        return chunks
    
    def _create_chunk(self, content: str, file_path: Path, chunk_index: int) -> DocumentChunk:
        """청크 객체 생성"""
        # 표 데이터 여부 확인 (더 정확한 패턴)
        is_table_data = "표 제목:" in content and "표 구조:" in content

        metadata = {
            'is_table_data': is_table_data,
            'content_type': 'table' if is_table_data else 'text'
        }

        # 표 데이터인 경우 추가 메타데이터
        if is_table_data:
            # 표 제목 추출
            lines = content.split('\n')
            for line in lines:
                if line.startswith("표 제목:"):
                    metadata['table_title'] = line.replace("표 제목:", "").strip()
                    break

            # 표 구조 정보 추출
            for line in lines:
                if "총 행 수:" in line:
                    import re
                    match = re.search(r'총 행 수: (\d+)행', line)
                    if match:
                        metadata['table_rows'] = int(match.group(1))
                    break

        return self._create_document_chunk(
            content=content,
            file_path=file_path,
            chunk_index=chunk_index,
            metadata=metadata,
            chunk_type='table' if is_table_data else 'text',
        )
    
    def process_directory(self, directory_path: Union[str, Path], force_process: bool = False) -> List[DocumentChunk]:
        """디렉토리 내 모든 파일 처리"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"디렉토리를 찾을 수 없습니다: {directory_path}")
        
        all_chunks = []
        processed_files = 0
        skipped_files = 0
        
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    chunks = self.process_file(file_path, force_process)
                    if chunks:
                        all_chunks.extend(chunks)
                        processed_files += 1
                    else:
                        skipped_files += 1
                except Exception as e:
                    self.logger.error(f"파일 처리 실패: {file_path}, 오류: {str(e)}")
                    continue
        
        self.logger.info(f"디렉토리 처리 완료: {directory_path}, 처리된 파일: {processed_files}, 건너뛴 파일: {skipped_files}, 총 청크 수: {len(all_chunks)}")
        return all_chunks
    
    def save_processed_chunks(self, chunks: List[DocumentChunk], output_dir: Union[str, Path]):
        """처리된 청크를 파일로 저장"""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        for chunk in chunks:
            output_file = output_dir / f"{chunk.chunk_id}.json"
            
            chunk_data = {
                'content': chunk.content,
                'metadata': chunk.metadata,
                'chunk_id': chunk.chunk_id,
                'source_file': chunk.source_file,
                'chunk_index': chunk.chunk_index
            }
            
            import json
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(chunk_data, f, ensure_ascii=False, indent=2)
        
        self.logger.info(f"청크 저장 완료: {output_dir}, 파일 수: {len(chunks)}")
    
    def _process_docx_with_semantic_chunking(self, file_path: Path) -> List[DocumentChunk]:
        """
        DOCX 파일을 의미 기반 청킹으로 처리
        
        Args:
            file_path: DOCX 파일 경로
            
        Returns:
            의미 기반 청킹된 DocumentChunk 리스트
        """
        try:
            self.logger.info(f"의미 기반 청킹으로 DOCX 파일 처리: {file_path}")
            
            # 1. DOCX에서 문단 추출
            doc = Document(file_path)
            paragraphs = []
            current_h1 = None
            current_h2 = None
            current_h3 = None
            
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    # 제목 스타일 확인
                    style_name = para.style.name if para.style else ''
                    name_lower = style_name.lower()
                    
                    if 'heading 1' in name_lower or '제목 1' in style_name:
                        current_h1 = para.text.strip()
                        current_h2 = None
                        current_h3 = None
                    elif 'heading 2' in name_lower or '제목 2' in style_name:
                        current_h2 = para.text.strip()
                        current_h3 = None
                    elif 'heading 3' in name_lower or '제목 3' in style_name:
                        current_h3 = para.text.strip()
                    
                    paragraphs.append({
                        'text': para.text.strip(),
                        'heading': current_h1,
                        'sub_heading': current_h2,
                        'sub_sub_heading': current_h3
                    })
            
            # 2. 문장 단위 분할
            from src.modules.semantic_chunker import SemanticChunker
            from src.modules.embedding_module import EmbeddingManager
            
            all_sentences = []
            sentence_metadata_list = []
            
            semantic_chunker = SemanticChunker(
                embedding_manager=EmbeddingManager(),
                coherence_threshold=getattr(self.config, 'semantic_coherence_threshold', 0.7),
                max_chunk_size=self.config.chunk_size,
                min_chunk_size=getattr(self.config, 'min_chunk_size', 50),
                overlap_sentences=getattr(self.config, 'overlap_sentences', 3)
            )
            
            for para in paragraphs:
                sentences = semantic_chunker._split_into_sentences(para['text'])
                all_sentences.extend(sentences)
                
                # 각 문장의 메타데이터
                for sent in sentences:
                    sentence_metadata_list.append({
                        'heading': para['heading'],
                        'sub_heading': para['sub_heading'],
                        'sub_sub_heading': para['sub_sub_heading'],
                        'source_file': str(file_path)
                    })
            
            if len(all_sentences) == 0:
                self.logger.warning(f"문장을 추출할 수 없습니다: {file_path}")
                return []
            
            # 3. 의미 기반 청킹
            base_metadata = {
                'file_name': file_path.name,
                'file_path': str(file_path),
                'source_file': str(file_path),
                'file_size': file_path.stat().st_size,
                'chunking_method': 'semantic_coherence'
            }
            
            semantic_chunks = semantic_chunker.chunk_by_semantic_coherence(
                sentences=all_sentences,
                metadata=base_metadata
            )
            
            # 4. Sliding Window 오버랩 적용
            overlapping_contents = semantic_chunker.create_overlapping_chunks(
                semantic_chunks,
                overlap_sentences=getattr(self.config, 'overlap_sentences', 3)
            )
            
            # 5. DocumentChunk 객체 생성
            document_chunks = []
            for i, (semantic_chunk, overlapping_content) in enumerate(zip(semantic_chunks, overlapping_contents)):
                # 문장 메타데이터 병합
                chunk_metadata = {
                    **base_metadata,
                    'chunk_size': len(overlapping_content),
                    'chunk_index': i,
                    'coherence_score': semantic_chunk.coherence_score,
                    'has_overlap': i > 0,
                    'start_sentence_idx': semantic_chunk.start_sentence_idx,
                    'end_sentence_idx': semantic_chunk.end_sentence_idx,
                    'heading': sentence_metadata_list[semantic_chunk.start_sentence_idx]['heading'] if semantic_chunk.start_sentence_idx < len(sentence_metadata_list) else None,
                    'sub-heading': sentence_metadata_list[semantic_chunk.start_sentence_idx]['sub_heading'] if semantic_chunk.start_sentence_idx < len(sentence_metadata_list) else None,
                    'sub-sub-heading': sentence_metadata_list[semantic_chunk.start_sentence_idx]['sub_sub_heading'] if semantic_chunk.start_sentence_idx < len(sentence_metadata_list) else None
                }
                
                heading_path = [
                    chunk_metadata.get('heading'),
                    chunk_metadata.get('sub-heading'),
                    chunk_metadata.get('sub-sub-heading'),
                ]

                chunk = self._create_document_chunk(
                    content=overlapping_content,
                    file_path=file_path,
                    chunk_index=i,
                    metadata=chunk_metadata,
                    heading_path=heading_path,
                    chunk_type='semantic',
                    chunk_id_override=f"{self.current_document_metadata.doc_id}_semantic_{i}" if self.current_document_metadata else None,
                )
                document_chunks.append(chunk)
            
            self.logger.info(f"의미 기반 청킹 완료: {file_path}, 청크 수: {len(document_chunks)}")
            return document_chunks
            
        except Exception as e:
            self.logger.error(f"의미 기반 청킹 실패: {file_path}, 오류: {str(e)}")
            import traceback
            self.logger.error(f"상세 오류: {traceback.format_exc()}")
            # 폴백: 기본 청킹 방식 사용
            self.logger.info("기본 청킹 방식으로 폴백")
            content = self._process_docx(file_path)
            return self._chunk_content(content, file_path)


def process_documents(input_dir: str, output_dir: str) -> List[DocumentChunk]:
    """문서 처리 함수"""
    processor = DocumentProcessor()
    chunks = processor.process_directory(input_dir)
    processor.save_processed_chunks(chunks, output_dir)
    return chunks
